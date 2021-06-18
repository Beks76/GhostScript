from docx.api import Document

import re
import sys 

class Parser:
    def __init__(self, url: str, to_url: str) -> None:
        self.url = url
        self.destination = to_url
    
    def get_kaspi_id(self):
        document = Document(self.url)
        ids_arr = []
        for p in range(len(document.tables)):
            table = document.tables[p]
            ids = []
            for i, row in enumerate(table.rows):
                print(i)
                if i != 0:
                    ids.append([cell.text for cell in row.cells][3])
                    s = ""
                    digits = re.findall(r'\d', ids[i-1])
                    ids_arr.append(s.join(digits))

        self.arr_to_file(ids_arr, self.destination)

    def arr_to_file(self, source, destination):
        with open(destination, 'w') as f:
            for item in source:
                f.write("%s\n" % item)

    

if __name__ == '__main__':
    url = str(sys.argv[1])
    to_url = str(sys.argv[2])
    parse = Parser(url, to_url)
    parse.get_kaspi_id()
    